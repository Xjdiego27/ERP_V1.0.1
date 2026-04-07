# ============================================
# RUTAS PLANTILLAS — Generación de documentos
# a partir de plantillas DOCX con placeholders
# que se llenan con datos de la BD y/o selección
# ============================================

import os
import re
from io import BytesIO
from datetime import datetime
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse, FileResponse
from sqlalchemy.orm import Session
from pydantic import BaseModel
from typing import Dict, Optional
from docx import Document as DocxDocument

from database import (
    get_db, Personal, Contrato, Distrito, Area, Cargo
)
from auth_token import verificar_token

router = APIRouter()

# Carpeta de plantillas DOCX
TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "templates")

# Carpeta raíz para documentos generados (organizados por apellido)
DOCS_GENERADOS_DIR = os.path.join(os.path.dirname(__file__), "documentos_generados")
os.makedirs(DOCS_GENERADOS_DIR, exist_ok=True)


# ─── HELPERS ────────────────────────────────────

def _numero_a_texto(n):
    """Convierte un número a su representación en texto en español."""
    unidades = ['', 'uno', 'dos', 'tres', 'cuatro', 'cinco', 'seis', 'siete', 'ocho', 'nueve']
    decenas = ['', 'diez', 'veinte', 'treinta', 'cuarenta', 'cincuenta',
               'sesenta', 'setenta', 'ochenta', 'noventa']
    especiales = {
        11: 'once', 12: 'doce', 13: 'trece', 14: 'catorce', 15: 'quince',
        16: 'dieciséis', 17: 'diecisiete', 18: 'dieciocho', 19: 'diecinueve',
        21: 'veintiuno', 22: 'veintidós', 23: 'veintitrés', 24: 'veinticuatro',
        25: 'veinticinco', 26: 'veintiséis', 27: 'veintisiete', 28: 'veintiocho',
        29: 'veintinueve',
    }
    centenas = ['', 'ciento', 'doscientos', 'trescientos', 'cuatrocientos',
                'quinientos', 'seiscientos', 'setecientos', 'ochocientos', 'novecientos']

    if n == 0:
        return 'cero'
    if n == 100:
        return 'cien'

    resultado = ''

    if n >= 1000:
        miles = n // 1000
        resto = n % 1000
        if miles == 1:
            resultado = 'mil'
        else:
            resultado = _numero_a_texto(miles) + ' mil'
        if resto > 0:
            resultado += ' ' + _numero_a_texto(resto)
        return resultado.strip()

    if n >= 100:
        resultado = centenas[n // 100]
        resto = n % 100
        if resto > 0:
            resultado += ' ' + _numero_a_texto(resto)
        return resultado.strip()

    if n in especiales:
        return especiales[n]

    if n >= 10:
        d = n // 10
        u = n % 10
        if u == 0:
            return decenas[d]
        return decenas[d] + ' y ' + unidades[u]

    return unidades[n]


def _sueldo_a_texto(sueldo_str):
    """Convierte un sueldo como '2500.00' a 'dos mil quinientos con 00/100 soles'."""
    try:
        monto = float(str(sueldo_str).replace(',', ''))
        entero = int(monto)
        decimales = int(round((monto - entero) * 100))
        texto = _numero_a_texto(entero)
        return f"{texto} con {decimales:02d}/100 soles"
    except (ValueError, TypeError):
        return str(sueldo_str) if sueldo_str else ''


MESES_ES = {
    1: 'enero', 2: 'febrero', 3: 'marzo', 4: 'abril',
    5: 'mayo', 6: 'junio', 7: 'julio', 8: 'agosto',
    9: 'septiembre', 10: 'octubre', 11: 'noviembre', 12: 'diciembre'
}


def _obtener_datos_auto(id_personal: int, db: Session):
    """Extrae datos automáticos del empleado desde la BD."""
    persona = db.query(Personal).filter(Personal.ID_PERSONAL == id_personal).first()
    if not persona:
        return {}

    contrato = db.query(Contrato).filter(
        Contrato.ID_PERSONAL == id_personal,
        Contrato.ID_ESTADO_CONTRATO == 1
    ).first()

    # Distrito
    distrito_nombre = ''
    id_distr = getattr(persona, 'ID_DISTR', None)
    if id_distr and Distrito:
        dist = db.query(Distrito).filter(Distrito.ID_DISTR == id_distr).first()
        if dist:
            distrito_nombre = dist.DESCRIP if hasattr(dist, 'DESCRIP') else ''

    # Departamento y Provincia (resuelto via FK ID_PROV_DEPART)
    depart_y_provinc = ''
    id_dep = getattr(persona, 'ID_PROV_DEPART', None)
    if id_dep:
        try:
            from sqlalchemy import text as _text
            row = db.execute(_text("SELECT DESCRIP FROM provincia_departamento WHERE ID_PROV_DEPART = :id"), {"id": id_dep}).first()
            if row:
                depart_y_provinc = row[0]
        except Exception:
            pass

    # Dirección
    direccion = getattr(persona, 'DIRECCION', '') or ''

    # Sueldo del contrato vigente
    sueldo = ''
    if contrato and hasattr(contrato, 'SUELDO'):
        sueldo = str(contrato.SUELDO) if contrato.SUELDO else ''

    # Cargo del contrato vigente
    cargo_nombre = ''
    if contrato and hasattr(contrato, 'ID_CARGO') and contrato.ID_CARGO and Cargo:
        cargo_obj = db.query(Cargo).filter(Cargo.ID_CARGO == contrato.ID_CARGO).first()
        if cargo_obj:
            cargo_nombre = cargo_obj.DESCRIP if hasattr(cargo_obj, 'DESCRIP') else ''

    # Fecha fin contrato
    fecha_fin_contrato = ''
    mes_fin_contrato = ''
    anio_fin_contrato = ''
    if contrato and hasattr(contrato, 'FECH_CESE') and contrato.FECH_CESE:
        fc = contrato.FECH_CESE
        fecha_fin_contrato = str(fc.day) if hasattr(fc, 'day') else str(fc)
        mes_fin_contrato = MESES_ES.get(fc.month, '') if hasattr(fc, 'month') else ''
        anio_fin_contrato = str(fc.year) if hasattr(fc, 'year') else ''

    hoy = datetime.now()

    datos = {
        'nombres': persona.NOMBRES if persona.NOMBRES else '',
        'ape_paterno': persona.APE_PATERNO if persona.APE_PATERNO else '',
        'ape_materno': persona.APE_MATERNO if persona.APE_MATERNO else '',
        'num_doc': persona.NUM_DOC if persona.NUM_DOC else '',
        'direccion': direccion,
        'distrito': distrito_nombre,
        'depart_y_provinc': depart_y_provinc,
        'cargo': cargo_nombre,
        # Fecha de generación
        'dia que se genera': str(hoy.day),
        'mes que se genera': MESES_ES.get(hoy.month, ''),
        'año que se genera': str(hoy.year),
        # Sueldo
        'sueldo (en numeros)': sueldo,
        'sueldo en texto': _sueldo_a_texto(sueldo) if sueldo else '',
        # Fin de contrato
        'fecha_fin_contrato': fecha_fin_contrato,
        'mes_fin_contrato': mes_fin_contrato,
        'año_fin_contrato': anio_fin_contrato,
        # Fecha compuesta de generación (ej: "27 de marzo de 2026")
        'dia_mes_año_que se genera': f"{hoy.day} de {MESES_ES.get(hoy.month, '')} de {hoy.year}",
        # Alias (algunas plantillas usan nombre diferente)
        'provincia_y_departamento': depart_y_provinc,
    }
    return datos


def _extraer_placeholders(doc_path: str):
    """Extrae todos los placeholders {campo} de un documento DOCX."""
    doc = DocxDocument(doc_path)
    placeholders = set()

    for para in doc.paragraphs:
        texto = para.text
        matches = re.findall(r'\{([^}]+)\}', texto)
        placeholders.update(matches)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    matches = re.findall(r'\{([^}]+)\}', para.text)
                    placeholders.update(matches)

    return list(placeholders)


# Campos que se llenan automáticamente desde la BD
CAMPOS_AUTO = {
    'nombres', 'ape_paterno', 'ape_materno', 'num_doc',
    'direccion', 'distrito', 'depart_y_provinc', 'provincia_y_departamento',
    'cargo',
    'dia que se genera', 'mes que se genera', 'año que se genera',
    'dia_mes_año_que se genera',
    'sueldo (en numeros)', 'sueldo en texto',
    'fecha_fin_contrato', 'mes_fin_contrato', 'año_fin_contrato',
}


# Versión en minúsculas para comparación insensible a mayúsculas
CAMPOS_AUTO_LOWER = {c.lower() for c in CAMPOS_AUTO}


def _clasificar_campos(placeholders):
    """Clasifica cada placeholder como 'auto' o 'manual' (insensible a mayúsculas)."""
    resultado = []
    for campo in sorted(placeholders):
        tipo = 'auto' if campo.lower() in CAMPOS_AUTO_LOWER else 'manual'
        resultado.append({'campo': campo, 'tipo': tipo})
    return resultado


def _reemplazar_en_parrafo(para, datos):
    """Reemplaza placeholders en un párrafo preservando formato (insensible a mayúsculas)."""
    texto_completo = para.text
    if '{' not in texto_completo:
        return

    # Construir mapa insensible a mayúsculas
    datos_lower = {k.lower(): v for k, v in datos.items()}

    # Buscar todos los placeholders en el texto y reemplazar
    matches = re.findall(r'\{([^}]+)\}', texto_completo)
    for match in matches:
        val = datos_lower.get(match.lower())
        if val is not None:
            texto_completo = texto_completo.replace('{' + match + '}', str(val))

    # Reconstruir runs preservando formato del primer run
    if para.runs:
        for i, run in enumerate(para.runs):
            if i == 0:
                run.text = texto_completo
            else:
                run.text = ''
    else:
        para.text = texto_completo


def _rellenar_documento(doc_path: str, datos: dict):
    """Abre el DOCX, reemplaza placeholders y retorna el documento."""
    doc = DocxDocument(doc_path)

    for para in doc.paragraphs:
        _reemplazar_en_parrafo(para, datos)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _reemplazar_en_parrafo(para, datos)

    return doc


# ─── SCHEMAS ────────────────────────────────────

class GenerarDocumentoRequest(BaseModel):
    campos_manuales: Dict[str, str] = {}


# ─── ENDPOINTS ──────────────────────────────────

def _resolver_carpeta_plantillas(id_emp=None, id_depart=None):
    """Resuelve la carpeta de plantillas según empresa y departamento.
    Busca en orden: templates/{id_emp}/{id_depart}/ → templates/{id_emp}/ → templates/
    """
    if id_emp and id_depart:
        ruta = os.path.join(TEMPLATES_DIR, str(id_emp), str(id_depart))
        if os.path.isdir(ruta):
            return ruta
    if id_emp:
        ruta = os.path.join(TEMPLATES_DIR, str(id_emp))
        if os.path.isdir(ruta):
            return ruta
    return TEMPLATES_DIR


def _listar_docx(carpeta):
    """Lista archivos .docx en una carpeta (sin subcarpetas)."""
    if not os.path.isdir(carpeta):
        return []
    archivos = []
    for f in os.listdir(carpeta):
        if f.lower().endswith('.docx') and not f.startswith('~$'):
            nombre_sin_ext = os.path.splitext(f)[0]
            archivos.append({
                'archivo': f,
                'nombre': nombre_sin_ext,
            })
    archivos.sort(key=lambda x: x['nombre'])
    return archivos


@router.get("/plantillas")
def listar_plantillas(
    id_depart: Optional[int] = Query(None),
    token: dict = Depends(verificar_token),
):
    """Lista las plantillas DOCX disponibles, filtradas por empresa (del token) y departamento."""
    id_emp = token.get("id_emp")
    carpeta = _resolver_carpeta_plantillas(id_emp, id_depart)
    return _listar_docx(carpeta)


@router.get("/plantillas/{nombre_archivo}/campos")
def obtener_campos_plantilla(
    nombre_archivo: str,
    id_personal: int = Query(...),
    id_depart: Optional[int] = Query(None),
    db: Session = Depends(get_db),
    token: dict = Depends(verificar_token)
):
    """Analiza una plantilla y retorna sus campos clasificados como auto/manual,
    incluyendo los valores automaticos ya resueltos desde la BD."""

    id_emp = token.get("id_emp")
    carpeta = _resolver_carpeta_plantillas(id_emp, id_depart)
    ruta = os.path.join(carpeta, nombre_archivo)
    if not os.path.isfile(ruta):
        raise HTTPException(status_code=404, detail="Plantilla no encontrada")

    placeholders = _extraer_placeholders(ruta)
    campos = _clasificar_campos(placeholders)
    datos_auto = _obtener_datos_auto(id_personal, db)

    # Agregar valores (búsqueda insensible a mayúsculas)
    datos_auto_lower = {k.lower(): v for k, v in datos_auto.items()}
    for c in campos:
        if c['tipo'] == 'auto':
            c['valor'] = datos_auto_lower.get(c['campo'].lower(), '')
        else:
            # Para campos manuales, sugerir valor desde BD si existe
            sugerido = datos_auto_lower.get(c['campo'].lower(), '')
            if sugerido:
                c['valor_sugerido'] = sugerido

    return {
        'archivo': nombre_archivo,
        'campos': campos,
    }


@router.post("/plantillas/{nombre_archivo}/generar")
def generar_documento(
    nombre_archivo: str,
    data: GenerarDocumentoRequest,
    id_personal: int = Query(...),
    formato: str = Query('docx', pattern='^(docx|pdf)$'),
    id_depart: Optional[int] = Query(None),
    db: Session = Depends(get_db),
    token: dict = Depends(verificar_token)
):
    """Genera el documento rellenado con datos auto + manuales.
    Retorna el archivo DOCX (o PDF si se especifica formato=pdf)."""

    id_emp = token.get("id_emp")
    carpeta = _resolver_carpeta_plantillas(id_emp, id_depart)
    ruta = os.path.join(carpeta, nombre_archivo)
    if not os.path.isfile(ruta):
        raise HTTPException(status_code=404, detail="Plantilla no encontrada")

    # Datos automáticos de la BD
    datos_auto = _obtener_datos_auto(id_personal, db)

    # Combinar: auto + manuales (los manuales sobreescriben si coinciden)
    datos_final = {**datos_auto, **data.campos_manuales}

    # Rellenar el documento
    doc = _rellenar_documento(ruta, datos_final)

    # Guardar en memoria
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    nombre_descarga = os.path.splitext(nombre_archivo)[0]

    # ── Guardar copia en carpeta del empleado (por apellidos) ──
    _guardar_documento_empleado(id_personal, nombre_descarga, formato, buffer, db)
    buffer.seek(0)  # re-posicionar después de guardar

    if formato == 'pdf':
        try:
            import tempfile
            import uuid as _uuid
            import pythoncom
            import win32com.client

            tmp_dir = tempfile.gettempdir()
            uid = _uuid.uuid4().hex[:8]
            tmp_docx = os.path.join(tmp_dir, f"doc_{uid}_{nombre_descarga}.docx")
            tmp_pdf  = os.path.join(tmp_dir, f"doc_{uid}_{nombre_descarga}.pdf")

            # Guardar DOCX temporal
            with open(tmp_docx, 'wb') as f:
                f.write(buffer.getvalue())

            # Convertir usando COM directamente con CoInitialize
            pythoncom.CoInitialize()
            try:
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                word.DisplayAlerts = False
                doc_w = word.Documents.Open(os.path.abspath(tmp_docx))
                doc_w.SaveAs(os.path.abspath(tmp_pdf), FileFormat=17)  # 17 = wdFormatPDF
                doc_w.Close(False)
                word.Quit()
            finally:
                pythoncom.CoUninitialize()

            if os.path.isfile(tmp_pdf):
                with open(tmp_pdf, 'rb') as f:
                    pdf_bytes = f.read()

                # Guardar copia PDF en carpeta del empleado
                _guardar_documento_empleado(id_personal, nombre_descarga, 'pdf', BytesIO(pdf_bytes), db)

                # Limpiar temporales
                try:
                    os.remove(tmp_docx)
                    os.remove(tmp_pdf)
                except OSError:
                    pass

                return StreamingResponse(
                    BytesIO(pdf_bytes),
                    media_type="application/pdf",
                    headers={
                        "Content-Disposition": f'attachment; filename="{nombre_descarga}.pdf"'
                    }
                )
            else:
                raise HTTPException(status_code=500, detail="Error al convertir a PDF")

        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error al generar PDF: {str(e)}")

    # DOCX por defecto
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{nombre_descarga}.docx"'
        }
    )


# ═══════════════════════════════════════════════════════
# GUARDAR DOCUMENTOS EN CARPETA POR APELLIDO DEL EMPLEADO
# ═══════════════════════════════════════════════════════

def _normalizar_nombre_carpeta(texto: str) -> str:
    """Normaliza un texto para usarlo como nombre de carpeta seguro."""
    import unicodedata
    # Quitar acentos
    nfkd = unicodedata.normalize('NFKD', texto)
    sin_acentos = ''.join(c for c in nfkd if not unicodedata.combining(c))
    # Solo alfanuméricos, espacios y guiones
    limpio = re.sub(r'[^\w\s-]', '', sin_acentos).strip()
    # Reemplazar espacios por guión bajo
    limpio = re.sub(r'\s+', '_', limpio)
    return limpio.upper() or 'SIN_NOMBRE'


def _obtener_carpeta_empleado(id_personal: int, db: Session) -> str:
    """Retorna la ruta de la carpeta del empleado organizada por apellidos.
    Formato: DOCS_GENERADOS_DIR / APE_PATERNO_APE_MATERNO_NOMBRES /
    Si el empleado ya tiene una carpeta (puede que cambiaron de nombre), la reutiliza."""

    persona = db.query(Personal).filter(Personal.ID_PERSONAL == id_personal).first()
    if not persona:
        return os.path.join(DOCS_GENERADOS_DIR, f"empleado_{id_personal}")

    ape_pat = (persona.APE_PATERNO or '').strip()
    ape_mat = (persona.APE_MATERNO or '').strip()
    nombres = (persona.NOMBRES or '').strip()

    nombre_carpeta = _normalizar_nombre_carpeta(f"{ape_pat} {ape_mat}, {nombres}")
    carpeta = os.path.join(DOCS_GENERADOS_DIR, nombre_carpeta)
    os.makedirs(carpeta, exist_ok=True)
    return carpeta


def _guardar_documento_empleado(id_personal: int, nombre_doc: str, formato: str, buffer: BytesIO, db: Session):
    """Guarda una copia del documento generado en la carpeta del empleado.
    Si ya existe un archivo con el mismo nombre, lo sobrescribe (actualiza)."""
    try:
        carpeta = _obtener_carpeta_empleado(id_personal, db)
        # Nombre con timestamp para evitar confusiones
        fecha_str = datetime.now().strftime('%Y%m%d_%H%M%S')
        nombre_archivo = f"{nombre_doc}_{fecha_str}.{formato}"
        ruta_destino = os.path.join(carpeta, nombre_archivo)

        pos_actual = buffer.tell()
        buffer.seek(0)
        with open(ruta_destino, 'wb') as f:
            f.write(buffer.read())
        buffer.seek(pos_actual)
    except Exception:
        # No fallar la generación si el guardado falla
        pass


# ═══════════════════════════════════════════════════════
# ENDPOINTS PARA DOCUMENTOS GENERADOS (guardados)
# ═══════════════════════════════════════════════════════

@router.get("/plantillas/generados/{id_personal}")
def listar_documentos_generados(
    id_personal: int,
    db: Session = Depends(get_db),
    token: dict = Depends(verificar_token)
):
    """Lista los documentos generados y guardados para un empleado."""
    carpeta = _obtener_carpeta_empleado(id_personal, db)

    if not os.path.isdir(carpeta):
        return []

    archivos = []
    for f in sorted(os.listdir(carpeta), reverse=True):
        ruta = os.path.join(carpeta, f)
        if not os.path.isfile(ruta):
            continue
        ext = os.path.splitext(f)[1].lower()
        if ext not in ('.docx', '.pdf'):
            continue
        stat = os.stat(ruta)
        archivos.append({
            'nombre': f,
            'formato': ext.replace('.', ''),
            'tamano': stat.st_size,
            'fecha': datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M'),
        })

    return archivos


@router.get("/plantillas/generados/{id_personal}/descargar")
def descargar_documento_generado(
    id_personal: int,
    archivo: str = Query(...),
    db: Session = Depends(get_db),
    token: dict = Depends(verificar_token)
):
    """Descarga un documento generado previamente guardado."""
    carpeta = _obtener_carpeta_empleado(id_personal, db)
    ruta = os.path.join(carpeta, archivo)

    # Seguridad: evitar path traversal
    ruta_real = os.path.realpath(ruta)
    carpeta_real = os.path.realpath(carpeta)
    if not ruta_real.startswith(carpeta_real):
        raise HTTPException(status_code=400, detail="Ruta no válida")

    if not os.path.isfile(ruta):
        raise HTTPException(status_code=404, detail="Archivo no encontrado")

    ext = os.path.splitext(archivo)[1].lower()
    media = "application/pdf" if ext == '.pdf' else \
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    return FileResponse(
        ruta,
        media_type=media,
        filename=archivo,
    )


@router.delete("/plantillas/generados/{id_personal}")
def eliminar_documento_generado(
    id_personal: int,
    archivo: str = Query(...),
    db: Session = Depends(get_db),
    token: dict = Depends(verificar_token)
):
    """Elimina un documento generado guardado."""
    carpeta = _obtener_carpeta_empleado(id_personal, db)
    ruta = os.path.join(carpeta, archivo)

    # Seguridad: evitar path traversal
    ruta_real = os.path.realpath(ruta)
    carpeta_real = os.path.realpath(carpeta)
    if not ruta_real.startswith(carpeta_real):
        raise HTTPException(status_code=400, detail="Ruta no válida")

    if not os.path.isfile(ruta):
        raise HTTPException(status_code=404, detail="Archivo no encontrado")

    os.remove(ruta)
    return {"ok": True, "mensaje": "Documento eliminado"}
